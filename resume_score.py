import os
import json

from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel
from pypdf import PdfReader
from docx import Document
load_dotenv()

my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("Invalid API key")

client = Groq(api_key = my_api_key)

model = "llama-3.3-70b-versatile"

class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferd_skill: list[str]
    minimum_experience: float | None
    education_requirement: list[str]
    responsibilities: list[str]
    
def parse_job_description(job_description):

    job_schema = JobD.model_json_schema()

    system_prompt = f"""
    You are an expert HR assistant.

    Your job is to analyze job descriptions and extract
    structured information from them.

    Return ONLY valid JSON matching this schema:

    {job_schema}

    IMPORTANT:
    Do NOT return the schema itself.
    Do NOT return fields like "properties", "title" or "type".
    Fill the schema with actual information extracted from
    the job description.

    If minimum experience is not mentioned, return null.
    If information for a list is missing, return an empty list.
    Do not invent information.
    """

    user_prompt = f"""
    Analyze the following job description:

    {job_description}
    """

    system_message = {
        "role": "system",
        "content": system_prompt
    }

    user_message = {
        "role": "user",
        "content": user_prompt
    }

    messages = [
        system_message,
        user_message
    ]

    response_format = {
        "type": "json_object"
    }

    response = client.chat.completions.create(
        model=model,
        messages=messages,
        response_format=response_format
    )

    answer = response.choices[0].message.content

    job_data = json.loads(answer)

    job = JobD(**job_data)

    return job

class MatchDetails(BaseModel):
    candidate_name: str | None = None
    matching_skills: list[str] = []
    missing_skills: list[str] = []
    experience_requirement_met: bool = False
    final_verdict: str = ""

class MatchResult(BaseModel):
    score: float
    details: MatchDetails
    
class Experience(BaseModel):
    company:str | None = None
    role:str | None = None
    duration:str | None = None
    description:str | None = None
    skill_used:list[str] = []
    
class Resume(BaseModel):
    name:str | None = None
    email:str | None = None
    phone:str | None = None
    total_experience_years:float | None = None
    experiences:list[Experience] = []
    skills:list[str] = []
    education:list[str] = []
    projects:list[str] = []
    certifications:list[str] = []


resume_schema = Resume.model_json_schema()

def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
        You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message= {
         "role":"user",
         "content": prompt
     }
    messages = [message]
    response_format = {
        "type":"json_object",
    }
    response = client.chat.completions.create(model=model,messages=messages,response_format=response_format,)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)

def parse_resume(resume_text):
    system_prompt = f"""
     You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f""" Parse the following resume:

    {resume_text}"""
    message_system = {
        "role":"system",
        "content": system_prompt
    }
    message_user = {
        "role":"user",
        "content": user_prompt
    }
    messages = [message_system, message_user]
    response_format = {
        "type":"json_object",
    }
    response = client.chat.completions.create(model=model,messages=messages,response_format=response_format,)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume

def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):  
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text

def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None
    
